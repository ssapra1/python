
# Ensure python-docx library is installed: pip install python-docx
from docx import Document

# Create a new Document with the updated information
updated_doc = Document()

# Add header with name and contact details
heading = updated_doc.add_heading('Saksham Sapra', level=1)
heading.style = 'Title'
updated_doc.add_paragraph("-" * 50)
updated_doc.add_paragraph(
    "Email: sapra.saksham1@gmail.com\n"
    "Phone: +91-8379074226\n"
    "LinkedIn: https://www.linkedin.com/in/saksham-sapra-253694148/\n"
    "GitHub:https://github.com/ssapra1/python\n"
    "Website: https://python-a4zrlhtjhzx8j9ktwopm5c.streamlit.app/"
)

# Add Professional Summary
updated_doc.add_heading('Professional Summary', level=2)
updated_doc.add_paragraph(
    "Dedicated Software Developer with 9+ years of experience in developing robust applications using Java, "
    "Spring Boot, Kafka, Microservices, and Cloud Technologies. Skilled in system integration, performance optimization, "
    "database development, and team leadership. Proven track record in delivering end-to-end solutions, maintaining large-scale systems, "
    "and improving application reliability."
)

# Add Skills Section
updated_doc.add_heading('Skills', level=2)
updated_doc.add_paragraph("**Programming Languages:** Java 8, Python")
updated_doc.add_paragraph("**Frameworks & Tools:** Streamlit, Spring Boot, Apache Flink, Multi-threading, Kafka, SOAP, REST")
updated_doc.add_paragraph("**Databases:** MongoDB")
updated_doc.add_paragraph("**Cloud Platforms:** Azure Cloud Containers")
updated_doc.add_paragraph("**Other Skills:** Microservices Architecture, Data Structures, Shell Scripting, Message Queues, Agile/Scrum")

# Add Professional Experience
updated_doc.add_heading('Professional Experience', level=2)

# Experience 1 - Citi Bank
updated_doc.add_heading('Senior Software Developer', level=3)
updated_doc.add_paragraph("Citi Bank | July 2021 - Present")
updated_doc.add_paragraph(
    "• Worked with the Helios system to develop and implement functionality for generating PDF reports. These reports consolidate detailed information on collateral, loans, and customer data, providing both enterprise and individual-level insights. Enabled customers to easily download these comprehensive reports for enhanced loan and collateral management.\n"
    "• Developed C1C Application to manage documents, batch processing, and validation systems using Java 8, Spring Boot Azure Cloud Containers, and MongoDB.\n"
    "• Enhanced Feed Processing Application to transfer retail loan data across systems, including NAS Mount, Shell Scripting, and server deployments.\n"
    "• Improved Kafka messaging systems with Avro schema transparency and developed a KPI Metrics library to integrate with Tableau Dashboards.\n"
    "• Led software architecture discussions, conducted code reviews, and provided mentorship to junior developers.\n"
    "• Delivered large-scale, secure systems for high-reliability e-commerce platforms."
)

# Experience 2 - Principal Financial Group
updated_doc.add_heading('Software Developer', level=3)
updated_doc.add_paragraph("Principal Financial Group | January 2018 - December 2019")
updated_doc.add_paragraph(
    "• Designed and developed call routing systems (IVR) to handle customer queries for insurance and retirement products.\n"
    "• Implemented fraud detection APIs and centralized authentication systems to streamline customer communication processes.\n"
    "• Built enterprise-level solutions to consolidate customer data into a unified web interface."
)

# Experience 3 - Optimus Project
updated_doc.add_heading('Senior Software Developer', level=3)
updated_doc.add_paragraph("Company ABC | December 2019 - July 2021")
updated_doc.add_paragraph(
    "• Worked on Optimus Project, resolving discrepancies in financial records using Java 8, Apache Flink, and Microservices.\n"
    "• Ensured high-quality reconciliation processes, enabling fraud detection and financial reporting."
)

# Add Education Section
updated_doc.add_heading('Education', level=2)
updated_doc.add_paragraph("• **B.Tech**, Amity University | June 2011 - May 2015")
updated_doc.add_paragraph("• KL Arya DAV Public School")

# Add Certifications Section
updated_doc.add_heading('Certifications', level=2)
updated_doc.add_paragraph(
    "• AWS Certified Developer - Associate\n"
)
# Add Achievements Section
updated_doc.add_heading('Achievements', level=2)
updated_doc.add_paragraph("• Achieved top 1% ranking globally in a competitive coding competition.")
updated_doc.add_paragraph("• Recognized as Employee of the Year for outstanding IT contributions.")
updated_doc.add_paragraph("• Played a major role in leading a cross-functional team to success.")

# Add Projects Section
updated_doc.add_heading('Projects', level=2)
updated_doc.add_paragraph(
    "• **KPI Metrics Dashboard**: Developed a Java-based library integrating Kafka and Tableau for real-time metrics visualization.\n"
    "• **Organizational Chatbot**: Implemented chatbot systems to automate policy-related queries within the organization."
)

# Add Languages Section
updated_doc.add_heading('Languages', level=2)
updated_doc.add_paragraph("• English")

# Add a cover page
updated_doc.add_page_break()
heading = updated_doc.add_heading("Resume - Saksham Sapra", level=1)
heading.style = "Title"
updated_doc.add_paragraph(
    "Prepared by: Saksham Sapra\n"
    "Position: Software Developer"
)
updated_doc.add_page_break()

# Save the updated document
updated_resume_path = "Updated_Resume_Saksham_Sapra_With_Cover_Page.docx"
updated_doc.save(updated_resume_path)

updated_resume_path
